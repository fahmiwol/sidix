"""
document_parser.py — SIDIX Document Parser
===========================================
Parser untuk dokumen Word, Excel, CSV, JSON, dan teks umum.
Semua pure Python / open-source library. Tidak pakai vendor API.

Research notes:
  - 318 cognitive expansion (input expansion)
"""
from __future__ import annotations

import csv
import io
import json
import os
from pathlib import Path
from typing import Any


def _ok(data: Any, note: str = "") -> dict:
    return {"ok": True, "data": data, "fallback_instructions": note, "citations": []}


def _fallback(instructions: str, data: Any = None) -> dict:
    return {"ok": False, "data": data, "fallback_instructions": instructions, "citations": []}


def parse_word(path: str) -> dict:
    """Ekstrak teks dari .docx (python-docx)."""
    if not os.path.exists(path):
        return _fallback(f"File tidak ditemukan: {path}")

    try:
        import docx  # type: ignore
        doc = docx.Document(path)
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        tables = []
        for table in doc.tables:
            rows = []
            for row in table.rows:
                rows.append([cell.text for cell in row.cells])
            tables.append(rows)

        return _ok({
            "backend": "python-docx",
            "extension": ".docx",
            "paragraphs": paragraphs,
            "paragraphs_count": len(paragraphs),
            "tables": tables,
            "tables_count": len(tables),
            "text": "\n".join(paragraphs),
        }, note="Untuk .doc (format lama), convert ke .docx dulu via libreoffice --headless.")
    except ImportError:
        return _fallback(
            "Library python-docx belum terpasang. Jalankan: pip install python-docx",
            data={"extension": ".docx"},
        )
    except Exception as exc:  # noqa: BLE001
        return _fallback(f"python-docx gagal: {exc}")


def parse_excel(path: str, sheet_index: int = 0) -> dict:
    """Ekstrak data dari .xlsx / .xls (openpyxl / xlrd)."""
    if not os.path.exists(path):
        return _fallback(f"File tidak ditemukan: {path}")

    errors = []

    # openpyxl untuk .xlsx
    if path.lower().endswith(".xlsx"):
        try:
            import openpyxl  # type: ignore
            wb = openpyxl.load_workbook(path, data_only=True)
            sheet = wb.worksheets[sheet_index] if sheet_index < len(wb.worksheets) else wb.active
            rows = []
            for row in sheet.iter_rows(values_only=True):
                rows.append([str(cell) if cell is not None else "" for cell in row])
            return _ok({
                "backend": "openpyxl",
                "extension": ".xlsx",
                "sheet_names": wb.sheetnames,
                "sheet_used": sheet.title,
                "rows": rows,
                "row_count": len(rows),
                "col_count": len(rows[0]) if rows else 0,
            })
        except ImportError:
            errors.append("openpyxl belum terpasang. Jalankan: pip install openpyxl")
        except Exception as exc:  # noqa: BLE001
            errors.append(f"openpyxl gagal: {exc}")

    # xlrd untuk .xls
    if path.lower().endswith(".xls"):
        try:
            import xlrd  # type: ignore
            wb = xlrd.open_workbook(path)
            sheet = wb.sheet_by_index(sheet_index) if sheet_index < wb.nsheets else wb.sheet_by_index(0)
            rows = []
            for r in range(sheet.nrows):
                rows.append([str(cell) if cell is not None else "" for cell in sheet.row_values(r)])
            return _ok({
                "backend": "xlrd",
                "extension": ".xls",
                "sheet_names": wb.sheet_names(),
                "sheet_used": sheet.name,
                "rows": rows,
                "row_count": len(rows),
                "col_count": len(rows[0]) if rows else 0,
            })
        except ImportError:
            errors.append("xlrd belum terpasang. Jalankan: pip install xlrd")
        except Exception as exc:  # noqa: BLE001
            errors.append(f"xlrd gagal: {exc}")

    return _fallback("; ".join(errors) if errors else "Format Excel tidak dikenal.")


def parse_csv(path: str, delimiter: str = ",") -> dict:
    """Ekstrak baris dari CSV / TSV."""
    if not os.path.exists(path):
        return _fallback(f"File tidak ditemukan: {path}")

    try:
        with open(path, "r", encoding="utf-8-sig", newline="") as f:
            reader = csv.reader(f, delimiter=delimiter)
            rows = [row for row in reader]
        return _ok({
            "backend": "csv",
            "extension": ".csv",
            "rows": rows,
            "row_count": len(rows),
            "col_count": len(rows[0]) if rows else 0,
        })
    except Exception as exc:  # noqa: BLE001
        return _fallback(f"CSV parse gagal: {exc}")


def parse_json(path: str) -> dict:
    """Load JSON file → Python object."""
    if not os.path.exists(path):
        return _fallback(f"File tidak ditemukan: {path}")

    try:
        with open(path, "r", encoding="utf-8") as f:
            data = json.load(f)
        return _ok({
            "backend": "json",
            "extension": ".json",
            "data": data,
            "type": type(data).__name__,
        })
    except json.JSONDecodeError as exc:
        return _fallback(f"JSON tidak valid: {exc}")
    except Exception as exc:  # noqa: BLE001
        return _fallback(f"JSON load gagal: {exc}")


def parse_text(path: str) -> dict:
    """Baca file teks biasa (.txt, .md, .py, dll)."""
    if not os.path.exists(path):
        return _fallback(f"File tidak ditemukan: {path}")

    try:
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            text = f.read()
        return _ok({
            "backend": "text",
            "extension": Path(path).suffix,
            "text": text,
            "char_count": len(text),
            "line_count": text.count("\n") + 1,
        })
    except Exception as exc:  # noqa: BLE001
        return _fallback(f"Text read gagal: {exc}")


def parse_document(path: str) -> dict:
    """Router otomatis berdasarkan ekstensi file."""
    ext = Path(path).suffix.lower()
    if ext in {".docx", ".doc"}:
        return parse_word(path)
    if ext in {".xlsx", ".xls"}:
        return parse_excel(path)
    if ext in {".csv", ".tsv"}:
        return parse_csv(path, delimiter="\t" if ext == ".tsv" else ",")
    if ext == ".json":
        return parse_json(path)
    if ext in {".txt", ".md", ".py", ".js", ".ts", ".html", ".css", ".yaml", ".yml", ".jsonl"}:
        return parse_text(path)

    return _fallback(
        f"Ekstensi '{ext}' belum didukung. "
        "Supported: .docx .xlsx .csv .json .txt .md .py .js .ts .html .css .yaml .jsonl",
        data={"extension": ext},
    )
